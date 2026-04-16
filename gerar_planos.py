#!/usr/bin/env python3
"""
Gerador Automático de Planos de Aula com IA (Google Gemini)
Lê o Escopo-Sequência Excel e gera planos de aula Word automaticamente.

Uso:
  py gerar_planos.py                      # Usa config.json
  py gerar_planos.py --listar             # Lista semanas disponíveis
  py gerar_planos.py --semanas 1,2,3      # Gera semanas específicas
  py gerar_planos.py --semanas 1-5        # Gera intervalo
  py gerar_planos.py --semanas todas      # Gera todas
  py gerar_planos.py --aba ADM --semanas 1
"""

import json
import shutil
import sys
import os
import re
import traceback
from pathlib import Path
from datetime import datetime, timedelta
import argparse


# --- Configuração ------------------------------------------------------------

def load_config(config_path="config.json"):
    path = Path(__file__).parent / config_path
    if not path.exists():
        print(f"ERRO: {config_path} não encontrado.")
        print("Execute com --criar-config para criar um modelo.")
        sys.exit(1)
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def criar_config_exemplo():
    exemplo = {
        "_ajuda": "Configure os campos abaixo e renomeie ou aponte para os arquivos corretos.",
        "escola": "Ayrton Senna",
        "professor": "Rafael Rodrigues Vieira",
        "ano_serie": "3B",
        "aba": "DADOS",
        "excel_path": "c:\\Users\\Desk\\Downloads\\Ed Profissional_Escopo-sequência_ANO2_2026 (1).xlsx",
        "template_path": "c:\\Users\\Desk\\Downloads\\3B - 02_03 _ 06_03 - Plano-de-Aula - Análise Exploratória....docx",
        "output_dir": "planos_gerados",
        "gemini_api_key": "SUA_CHAVE_AQUI",
        "data_inicio_semana1": "2026-02-03",
        "dias_aula": [0, 3],
        "_dias_aula_info": "0=Segunda, 1=Terça, 2=Quarta, 3=Quinta, 4=Sexta. Ex: [0,3] = Seg e Qui",
        "semanas": "todas"
    }
    path = Path(__file__).parent / "config.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(exemplo, f, ensure_ascii=False, indent=2)
    print(f"Config criado em: {path}")
    print("Edite o arquivo e defina sua chave do Gemini API.")


# --- Leitura do Excel ---------------------------------------------------------

def read_excel(excel_source, sheet_name, componente_filter=None, bimestre_filter=None):
    import openpyxl
    # excel_source pode ser um Path, string ou BytesIO (vido do Streamlit)
    wb = openpyxl.load_workbook(excel_source, data_only=True)

    if sheet_name not in wb.sheetnames:
        raise ValueError(f"Aba '{sheet_name}' não encontrada no arquivo Excel.")

    ws = wb[sheet_name]
    lessons = []

    for row in ws.iter_rows(min_row=2, values_only=True):
        if not row[1] or not row[10]:
            continue
        
        # Ignorar linhas de recomposição ou que não tenham número no bimestre
        try:
            bimestre_val = int(row[1])
        except (ValueError, TypeError):
            continue

        if row[0] and "recomposi" in str(row[0]).lower():
            continue
        
        # Filtro de Bimestre (se fornecido)
        if bimestre_filter and bimestre_val != int(bimestre_filter):
            continue

        # Lógica inteligente para identificar a matéria:
        comp_col = str(row[5] or "").strip()
        tech_col = str(row[3] or "").strip()
        
        if len(comp_col) > 100 and tech_col and len(tech_col) < 100:
            materia_real = tech_col
        else:
            materia_real = comp_col or tech_col

        if componente_filter and componente_filter.lower() not in materia_real.lower():
            continue

        lessons.append({
            "bimestre": bimestre_val,
            "natureza_aula": str(row[2] or "Teórica/Prática"), # Coluna 2 do Excel
            "competencia_tecnica": tech_col,
            "competencias_socioemocionais": str(row[4] or ""),
            "componente": materia_real,
            "qtd_aulas_semana": int(row[7] or 4),
            "unidade_curricular": str(row[8] or ""),
            "semana": int(row[10]),
            "tema_semana": str(row[11] or ""),
            "titulo_aula": str(row[12] or ""),
            "habilidades_tecnicas": str(row[13] or ""),
            "habilidades_socioemocionais": str(row[14] or ""),
            "objetivo": str(row[15] or ""),
        })

    return lessons


def get_bimestres(excel_source, sheet_name):
    """Retorna a lista de bimestres disponíveis na aba."""
    import openpyxl
    wb = openpyxl.load_workbook(excel_source, data_only=True)
    if sheet_name not in wb.sheetnames:
        return []
    ws = wb[sheet_name]
    bimestres = set()
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[1]:
            try:
                bimestres.add(int(row[1]))
            except (ValueError, TypeError):
                continue
    return sorted(list(bimestres))


def get_weeks(lessons):
    return sorted(set(l["semana"] for l in lessons))


def get_componentes(excel_source, sheet_name):
    """Retorna uma lista única e ordenada de componentes (matérias) na aba."""
    import openpyxl
    wb = openpyxl.load_workbook(excel_source, data_only=True)
    if sheet_name not in wb.sheetnames:
        return []
    ws = wb[sheet_name]
    componentes = set()
    for row in ws.iter_rows(min_row=2, values_only=True):
        comp_col = str(row[5] or "").strip()
        tech_col = str(row[3] or "").strip()
        
        if len(comp_col) > 100 and tech_col and len(tech_col) < 100:
            materia_real = tech_col
        else:
            materia_real = comp_col or tech_col
            
        if materia_real:
            componentes.add(materia_real)
    return sorted(list(componentes))


def get_week_lessons(lessons, week_num):
    return [l for l in lessons if l["semana"] == week_num]


def get_week_dates(start_date_str, week_num, days_of_week):
    # start_date_str é a data da SEMANA 1 (início do ano letivo)
    start = datetime.strptime(start_date_str, "%Y-%m-%d")
    
    # O deslocamento deve ser de (week_num - 1) semanas a partir da data inicial
    # Se a Semana 1 começa em 03/02, a Semana 8 começa 7 semanas depois.
    week_start = start + timedelta(weeks=week_num - 1)
    
    # Encontrar os dias específicos daquela semana
    # Ajustamos para o dia da semana (0=Segunda, etc.)
    # Primeiro dia útil daquela semana específica (baseado na data inicial)
    dates = [week_start + timedelta(days=d - days_of_week[0]) for d in days_of_week]
    
    first, last = min(dates), max(dates)
    return first.strftime("%d/%m/%Y"), last.strftime("%d/%m/%Y"), first, last


# --- Gemini API ---------------------------------------------------------------

def call_gemini(week_lessons, api_key):
    from google import genai
    from google.genai import types

    client = genai.Client(api_key=api_key)
    modelo = "gemini-2.0-flash"

    lesson = week_lessons[0]
    qtd = lesson["qtd_aulas_semana"]

    aulas_desc = "\n".join(
        f"  Aula {i+1}: {l['titulo_aula']}\n    Objetivo: {l['objetivo']}\n    Habilidades: {l['habilidades_tecnicas']}"
        for i, l in enumerate(week_lessons)
    )

    prompt = f"""Você é um Coordenador Pedagógico e Professor Expert em Educação Profissional Técnica.
Sua tarefa é elaborar o conteúdo para o componente: {lesson['componente']}.

REFERÊNCIA PEDAGÓGICA: BNCC, Metodologias Ativas (PBL, Cooperativa) e Contexto Profissional.
RECURSOS: Computadores, Lousa, Material Digital, Projetor.

DADOS DA SEMANA:
- Tema: {lesson['tema_semana']}
- Competência: {lesson['competencia_tecnica']}
- Aulas: {qtd} aulas de 50 min.

AULAS PARA DESENVOLVER:
{aulas_desc}

---
TAREFA: Gere o conteúdo dividido em 3 seções claras usando estas tags exatas:

<DESENVOLVIMENTO>
(Para cada aula, escreva Abertura, Desenvolvimento e Fechamento de forma detalhada e didática. Use apenas texto puro, sem markdown.)
</DESENVOLVIMENTO>

<AEE>
(Sugestões de adaptação para alunos com necessidades especiais (AEE) de forma integrada a este conteúdo técnico. Texto puro.)
</AEE>

<EXERCICIOS>
(Crie uma lista de 3 a 5 exercícios práticos ou teóricos com gabarito ao final. Texto puro.)
</EXERCICIOS>

REGRAS: 
- NÃO use markdown (sem **, sem #). 
- Seja específico e profissional.
- Mencione o uso dos recursos (computador, lousa, etc.).
"""

    response = client.models.generate_content(model=modelo, contents=prompt)
    return response.text


def gerar_desenvolvimento_fallback(week_lessons):
    """Fallback quando Gemini não está configurado. Retorna formato com tags."""
    desenv = "AULAS DA SEMANA:\n\n"
    for i, l in enumerate(week_lessons):
        desenv += f"AULA {i+1} - {l['titulo_aula']}\n"
        desenv += f"Objetivo: {l['objetivo']}\n"
        desenv += f"Habilidades: {l['habilidades_tecnicas']}\n\n"
    
    return f"""<DESENVOLVIMENTO>
{desenv}
</DESENVOLVIMENTO>
<AEE>
Adaptações sugeridas: Utilizar material visual impresso e tempo estendido para atividades práticas.
</AEE>
<EXERCICIOS>
1. Descreva o objetivo principal do tema desta semana.
2. Como este conteúdo se aplica no mercado de trabalho?
Gabarito: Respostas baseadas no escopo-sequência.
</EXERCICIOS>"""


# --- Manipulação do Word ------------------------------------------------------

def _unique_cells(table):
    """Yield unique cells using the underlying XML element as stable identity."""
    seen = set()
    for row in table.rows:
        for cell in row.cells:
            if cell._tc not in seen:
                seen.add(cell._tc)
                yield cell


def _find_cell(table, search_text):
    """Find first unique cell containing search_text."""
    for cell in _unique_cells(table):
        if search_text in cell.text:
            return cell
    return None


def _replace_para(para, new_text):
    """Replace paragraph text, preserving the first run's XML/formatting."""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def update_field(table, label, new_value, value_in_next_para=False):
    """Find cell containing label and update the value."""
    cell = _find_cell(table, label)
    if not cell:
        return False

    for i, para in enumerate(cell.paragraphs):
        if label in para.text:
            if value_in_next_para:
                if i + 1 < len(cell.paragraphs):
                    _replace_para(cell.paragraphs[i + 1], new_value)
            else:
                colon_pos = para.text.find(': ')
                if colon_pos >= 0:
                    prefix = para.text[:colon_pos + 2]
                else:
                    prefix = para.text.rstrip() + ' '
                _replace_para(para, prefix + new_value)
            return True
    return False


def add_desenvolvimento(table, texto):
    """Insere o texto de desenvolvimento na célula 'Desenvolvimento:'."""
    cell = _find_cell(table, "Desenvolvimento:")
    if not cell:
        return

    linhas = texto.strip().split('\n')
    for linha in linhas:
        cell.add_paragraph(linha.strip() if linha.strip() else "")


def fill_template(template_source, output_dest, week_lessons, config, first_date, last_date, week_num, conteudo_obj):
    from docx import Document
    from docx.enum.text import WD_BREAK
    import io

    doc = Document(template_source)
    table = doc.tables[0]

    lesson = week_lessons[0]
    qtd = lesson["qtd_aulas_semana"]
    aula_ini = (week_num - 1) * qtd + 1
    aula_fim = aula_ini + qtd - 1

    # Se conteudo_obj for string (formato antigo/fallback), transforma em dict
    if isinstance(conteudo_obj, str):
        conteudo_obj = extrair_secoes(conteudo_obj)

    # Preenchimento de campos básicos
    update_field(table, "ESCOLA: ", config["escola"])
    update_field(table, "PROFESSOR(A): ", config["professor"])
    update_field(table, "COMPONENTE CURRICULAR: ", lesson["componente"])
    update_field(table, "ANO/S", config["ano_serie"])
    update_field(table, "BIMESTRE: ", str(lesson["bimestre"]))
    
    aula_info = f"Aulas {aula_ini} a {aula_fim} - Semana {week_num} ({lesson['natureza_aula']})"
    update_field(table, "AULA NO ES: ", aula_info)
    
    update_field(table, "APRENDIZAGEM ESSENCIAL:", lesson["tema_semana"], value_in_next_para=True)
    update_field(table, "HABILIDADE RELACIONADA:", lesson["habilidades_tecnicas"], value_in_next_para=True)
    
    conhecimentos_previos = f"Conteúdos estudados anteriormente em {lesson['componente']}. Unidade: {lesson['unidade_curricular']}."
    update_field(table, "CONHECIMENTOS PR", conhecimentos_previos, value_in_next_para=True)
    update_field(table, "QUANT. DE AULAS PREVISTAS: ", str(qtd))
    
    objetivos = "\n".join(f"Aula {i+1}: {l['objetivo']}" for i, l in enumerate(week_lessons))
    update_field(table, "OBJETIVO DA AULA:", objetivos, value_in_next_para=True)
    update_field(table, "DATA DE ELABORA", datetime.now().strftime("%d/%m/%Y"), value_in_next_para=False)

    # Inserir Desenvolvimento e AEE na tabela
    add_desenvolvimento(table, conteudo_obj["desenvolvimento"])
    
    # Preencher Adaptação AEE se o campo existir
    cell_aee = _find_cell(table, "ADAPTAÇÃO AEE:")
    if cell_aee and conteudo_obj["aee"]:
        for linha in conteudo_obj["aee"].split('\n'):
            cell_aee.add_paragraph(linha.strip())

    # Adicionar Lista de Exercícios em uma nova página (Anexo)
    if conteudo_obj["exercicios"]:
        doc.add_page_break()
        doc.add_heading(f"ANEXO: LISTA DE EXERCÍCIOS - SEMANA {week_num}", level=1)
        doc.add_paragraph(f"Componente: {lesson['componente']}")
        doc.add_paragraph(f"Tema: {lesson['tema_semana']}\n")
        
        for linha in conteudo_obj["exercicios"].split('\n'):
            doc.add_paragraph(linha.strip())

    doc.save(output_dest)


# --- Processamento Central ---------------------------------------------------

def executar_geracao(config, excel_source, template_source, selected_weeks, output_callback=print):
    """
    Motor principal de geração. Pode ser chamado via CLI ou via Streamlit.
    excel_source e template_source podem ser caminhos ou BytesIO.
    Retorna uma lista de tuplas (filename, bytes_do_arquivo).
    """
    aba = config.get("aba")
    componente = config.get("componente", "")
    bimestre = config.get("bimestre")
    api_key = config.get("gemini_api_key", "")
    use_ai = api_key and api_key != "SUA_CHAVE_AQUI"

    lessons = read_excel(excel_source, aba, componente_filter=componente, bimestre_filter=bimestre)
    weeks_available = get_weeks(lessons)

    start_date = config.get("data_inicio_semana1", "2026-02-03")
    days_of_week = config.get("dias_aula", [0, 3])

    arquivos_gerados = [] # List of (filename, bytes)

    import io

    for week_num in selected_weeks:
        if week_num not in weeks_available:
            output_callback(f"AVISO: Semana {week_num} não encontrada no escopo. Pulando.")
            continue

        week_lessons = get_week_lessons(lessons, week_num)
        d1_str, d2_str, first_date, last_date = get_week_dates(start_date, week_num, days_of_week)

        output_callback(f"Processando Semana {week_num:02d} ({d1_str} a {d2_str})...")

        # Gerar conteúdo
        if use_ai:
            try:
                texto_ia = call_gemini(week_lessons, api_key)
                conteudo = extrair_secoes(texto_ia)
            except Exception as e:
                output_callback(f"  ERRO IA: {e}")
                conteudo = extrair_secoes(gerar_desenvolvimento_fallback(week_lessons))
        else:
            conteudo = extrair_secoes(gerar_desenvolvimento_fallback(week_lessons))

        # Criar documento em memória
        filename = gerar_nome_arquivo(config, week_lessons[0], first_date, last_date, week_num)
        output_stream = io.BytesIO()

        try:
            # Re-abre o template para cada plano para garantir que está limpo
            if hasattr(template_source, 'seek'):
                template_source.seek(0)

            fill_template(template_source, output_stream, week_lessons, config,
                         first_date, last_date, week_num, conteudo)

            arquivos_gerados.append((filename, output_stream.getvalue()))
        except Exception as e:
            output_callback(f"  ERRO ao preencher template: {e}")

    return arquivos_gerados


# --- Main ---------------------------------------------------------------------

def parse_semanas(semanas_str, weeks_available):
    if semanas_str == "todas":
        return weeks_available
    if "-" in semanas_str and "," not in semanas_str:
        partes = semanas_str.split("-")
        return list(range(int(partes[0]), int(partes[1]) + 1))
    return [int(s.strip()) for s in semanas_str.split(",")]


def main():
    parser = argparse.ArgumentParser(
        description="Gerador de Planos de Aula com IA",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("--aba", help="Aba do Excel (ex: DADOS, ADM, ENF)")
    parser.add_argument("--semanas", help="Semanas: '1,2,3' ou '1-5' ou 'todas'")
    parser.add_argument("--config", default="config.json")
    parser.add_argument("--listar", action="store_true", help="Lista semanas disponíveis")
    parser.add_argument("--criar-config", action="store_true", help="Cria config.json de exemplo")
    args = parser.parse_args()

    if args.criar_config:
        criar_config_exemplo()
        return

    config = load_config(args.config)
    aba = args.aba or config.get("aba")
    if not aba:
        print("ERRO: defina 'aba' no config.json ou use --aba")
        sys.exit(1)

    excel_path = config["excel_path"]
    template_path = config["template_path"]
    output_dir = Path(config.get("output_dir", "planos_gerados"))
    output_dir.mkdir(exist_ok=True)

    print(f"\n[Excel] Escopo-Sequência: {Path(excel_path).name}")
    print(f"   Aba selecionada: {aba}")
    componente = config.get("componente", None)
    lessons = read_excel(excel_path, aba, componente_filter=componente)
    weeks = get_weeks(lessons)
    sample = get_week_lessons(lessons, weeks[0])[0]
    print(f"   Componente: {sample['componente']}")
    print(f"   Semanas no escopo: {weeks[0]} a {weeks[-1]} ({len(weeks)} semanas, {len(lessons)} aulas)")

    if args.listar:
        print("\n[Lista] Semanas disponíveis:")
        for w in weeks:
            wl = get_week_lessons(lessons, w)
            if wl:
                temas = ", ".join(set(l["titulo_aula"].split(":")[0].strip() for l in wl))
                print(f"  Semana {w:3d} | Bimestre {wl[0]['bimestre']} | {wl[0]['tema_semana']}")
        return

    semanas_str = args.semanas or config.get("semanas", "todas")
    selected = parse_semanas(str(semanas_str), weeks)

    api_key = config.get("gemini_api_key", "")
    use_ai = api_key and api_key != "SUA_CHAVE_AQUI"

    if not use_ai:
        print("\nAVISO  Gemini não configurado. Planos serão gerados SEM conteúdo de IA.")
        print("   Configure 'gemini_api_key' no config.json")
        print("   Chave gratuita em: https://aistudio.google.com/app/apikey\n")

    start_date = config.get("data_inicio_semana1", "2026-02-03")
    days_of_week = config.get("dias_aula", [0, 3])

    print(f"[Data] Início semana 1: {start_date} | Dias de aula: {days_of_week}")
    print(f"[Gerando] Gerando {len(selected)} plano(s)...\n")

    # Executar a geração usando o novo motor
    arquivos = executar_geracao(config, excel_path, template_path, selected)

    # Salvar os arquivos em disco (comportamento CLI original)
    gerados = 0
    for filename, content in arquivos:
        output_path = output_dir / filename
        with open(output_path, "wb") as f:
            f.write(content)
        print(f"    OK {filename}")
        gerados += 1

    print(f"\n{'='*50}")
    print(f"Concluído: {gerados} plano(s) gerado(s)")
    print(f"[Pasta] Pasta de saída: {output_dir.absolute()}")


if __name__ == "__main__":
    main()
